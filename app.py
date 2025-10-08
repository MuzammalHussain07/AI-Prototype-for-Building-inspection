import streamlit as st
from infer import detect_cracks
from docx import Document
from io import BytesIO
import os

st.set_page_config(page_title="AI Crack Detection", layout="wide")

st.title("🏗️ Building Inspection Prototype — Crack Detection")

st.write("""
Upload inspection photos, and the system will automatically highlight **cracks and defects**.
You can also download a simple Word report.
""")

# ✅ Safe install fallback (ensures libs on Streamlit Cloud)
os.system("pip install --quiet ultralytics==8.2.50 opencv-python-headless==4.9.0.80")

# File uploader
uploaded_files = st.file_uploader(
    "Upload images (JPG/PNG)...",
    type=["jpg", "jpeg", "png"],
    accept_multiple_files=True
)

if uploaded_files:
    st.info(f"🖼️ {len(uploaded_files)} image(s) uploaded.")
    results_summary = []

    for uploaded_file in uploaded_files:
        image_bytes = uploaded_file.read()
        result_img, detections = detect_cracks(image_bytes)

        st.image(result_img, caption=f"Detection Result — {uploaded_file.name}", use_container_width=True)
        st.write("Detections:", detections if detections else "No cracks detected.")

        results_summary.append({
            "file": uploaded_file.name,
            "detections": detections if detections else ["No cracks detected"]
        })

    # Generate report button
    if st.button("📄 Generate Word Report"):
        doc = Document()
        doc.add_heading("Building Crack Detection Report", level=1)

        for item in results_summary:
            doc.add_heading(item["file"], level=2)
            for det in item["detections"]:
                doc.add_paragraph(f"- {det}")

        buf = BytesIO()
        doc.save(buf)
        buf.seek(0)
        st.download_button(
            label="⬇️ Download Word Report",
            data=buf,
            file_name="inspection_report.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        )
