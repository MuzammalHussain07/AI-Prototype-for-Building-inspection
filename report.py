# report.py
from docx import Document
from docx.shared import Inches

def create_report(output_path, image_path, detections):
    doc = Document()
    doc.add_heading("Building Inspection Report", level=1)
    doc.add_paragraph("Detected Defects: Crack Analysis")

    doc.add_picture(image_path, width=Inches(5))
    doc.add_paragraph("Detected cracks and their measurements:")

    table = doc.add_table(rows=1, cols=4)
    hdr = table.rows[0].cells
    hdr[0].text = "Crack Category"
    hdr[1].text = "Confidence"
    hdr[2].text = "Width (px)"
    hdr[3].text = "Width (mm)"

    for d in detections:
        row = table.add_row().cells
        row[0].text = str(d['category'])
        row[1].text = str(round(d['confidence'], 2))
        row[2].text = str(d['px_width'])
        row[3].text = str(d['mm_width']) if d['mm_width'] else "N/A"

    doc.save(output_path)
    return output_path
